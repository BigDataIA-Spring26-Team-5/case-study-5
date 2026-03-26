"""IC Memo Generator — produces an Investment Committee memo (Word .docx).

CS5 bonus expects a Word document. This implementation prefers `.docx` via
python-docx and falls back to PDF (weasyprint) or plain text if dependencies are
missing.
"""
from __future__ import annotations

import html
import logging
from datetime import datetime
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)

IC_MEMO_CSS = """\
body {
    font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
    margin: 40px 50px;
    font-size: 11pt;
    color: #333;
    line-height: 1.5;
}
h1 {
    font-size: 20pt;
    color: #1a365d;
    border-bottom: 2px solid #1a365d;
    padding-bottom: 8px;
    margin-bottom: 4px;
}
h2 {
    font-size: 14pt;
    color: #2c5282;
    margin-top: 28px;
    margin-bottom: 8px;
}
h3 {
    font-size: 12pt;
    color: #2c5282;
    margin-top: 20px;
    margin-bottom: 6px;
}
table {
    border-collapse: collapse;
    width: 100%;
    margin: 12px 0;
}
th {
    background-color: #1a365d;
    color: white;
    padding: 8px 12px;
    text-align: left;
    font-size: 10pt;
}
td {
    padding: 8px 12px;
    border: 1px solid #ccc;
    font-size: 10pt;
}
tr:nth-child(even) td {
    background-color: #f7fafc;
}
.meta {
    color: #666;
    font-size: 10pt;
    margin: 2px 0;
}
.confidential {
    color: #c53030;
    font-weight: bold;
    font-style: italic;
    border: 1px solid #c53030;
    padding: 6px 12px;
    display: inline-block;
    margin-top: 8px;
}
.recommendation {
    font-weight: bold;
    border-left: 4px solid;
    padding: 8px 14px;
    margin-top: 8px;
}
.rec-proceed {
    color: #276749;
    border-color: #276749;
    background-color: #f0fff4;
}
.rec-monitor {
    color: #975a16;
    border-color: #975a16;
    background-color: #fffff0;
}
.rec-conditional {
    color: #c53030;
    border-color: #c53030;
    background-color: #fff5f5;
}
ul.gaps {
    padding-left: 20px;
}
ul.gaps li {
    margin: 6px 0;
}
"""


class ICMemoGenerator:
    """Generates IC memo documents from DD results."""

    def _try_weasyprint(self) -> Optional[Any]:
        try:
            from weasyprint import HTML  # type: ignore[import]
            return HTML
        except Exception:
            return None

    def _try_docx(self) -> Optional[Any]:
        try:
            from docx import Document  # type: ignore[import]
            return Document
        except Exception:
            return None

    def generate(
        self,
        company_id: str,
        scoring_result: Dict[str, Any],
        gap_analysis: Dict[str, Any],
        ebitda_projection: Dict[str, Any],
        output_path: Optional[str] = None,
        output_format: str = "docx",
    ) -> str:
        """Generate IC memo and save as `.docx` (or PDF / `.txt` fallback).

        Returns:
            Path to the generated file.
        """
        org_air = scoring_result.get("org_air", 0.0)
        vr = scoring_result.get("vr_score", 0.0)
        hr = scoring_result.get("hr_score", 0.0)
        date_str = datetime.utcnow().strftime("%B %d, %Y")

        fmt = (output_format or "").lower().strip()
        if output_path:
            lower = output_path.lower()
            if lower.endswith(".docx"):
                fmt = "docx"
            elif lower.endswith(".pdf"):
                fmt = "pdf"
            elif lower.endswith(".txt"):
                fmt = "txt"

        if fmt == "docx":
            Document = self._try_docx()
            if Document is not None:
                try:
                    dim_scores = scoring_result.get("dimension_scores", {}) or {}
                    return self._generate_docx(
                        Document=Document,
                        company_id=company_id,
                        org_air=float(org_air or 0.0),
                        vr=float(vr or 0.0),
                        hr=float(hr or 0.0),
                        dimension_scores=dim_scores,
                        gap_analysis=gap_analysis,
                        ebitda_projection=ebitda_projection,
                        output_path=output_path,
                        date_str=date_str,
                    )
                except Exception as exc:
                    logger.exception("IC memo DOCX generation failed; falling back: %s", exc)
            else:
                logger.warning("python-docx not installed — falling back")

        if fmt == "txt":
            return self._generate_text(
                company_id, org_air, vr, hr, gap_analysis, ebitda_projection, output_path, date_str
            )

        HTML = self._try_weasyprint()
        if HTML is None:
            logger.warning("weasyprint not installed — saving as .txt")
            return self._generate_text(
                company_id, org_air, vr, hr, gap_analysis, ebitda_projection, output_path, date_str
            )

        try:
            cid = html.escape(str(company_id))

            # Executive summary
            readiness = (
                "Strong AI readiness positions for value creation."
                if org_air >= 70
                else "Improvement opportunities identified across key dimensions."
            )

            # Scoring table rows
            scoring_rows = ""
            for metric, score, bench in [
                ("Org-AI-R", org_air, 65.0),
                ("V^R (Valuation Readiness)", vr, 60.0),
                ("H^R (Human Capital Risk)", hr, 60.0),
            ]:
                scoring_rows += (
                    f"<tr><td>{html.escape(metric)}</td>"
                    f"<td>{float(score):.1f}</td>"
                    f"<td>{float(bench):.1f}</td></tr>\n"
                )

            # Dimension scores table
            dim_scores = scoring_result.get("dimension_scores", {})
            dim_html = ""
            if dim_scores:
                dim_rows = ""
                for dim, score in dim_scores.items():
                    dim_name = html.escape(str(dim).replace("_", " ").title())
                    dim_rows += f"<tr><td>{dim_name}</td><td>{float(score or 0.0):.1f}</td></tr>\n"
                dim_html = f"""
                <h3>Dimension Scores</h3>
                <table>
                    <thead><tr><th>Dimension</th><th>Score</th></tr></thead>
                    <tbody>{dim_rows}</tbody>
                </table>"""

            # Gap analysis
            gaps = gap_analysis.get("dimension_gaps", [])
            if gaps:
                gap_items = ""
                for gap in gaps[:5]:
                    dim_name = html.escape(str(gap.get("dimension", "")).replace("_", " ").title())
                    current = float(gap.get("current_score", 0) or 0)
                    target = float(gap.get("target_score", 0) or 0)
                    priority = html.escape(str(gap.get("priority", "N/A")))
                    gap_items += (
                        f"<li>{dim_name}: current {current:.1f} &rarr; "
                        f"target {target:.1f} [Priority: {priority}]</li>\n"
                    )
                gaps_html = f'<ul class="gaps">{gap_items}</ul>'
            else:
                gaps_html = "<p>No gap analysis data available.</p>"

            # EBITDA projection
            scenarios = ebitda_projection.get("scenarios", {})
            base_val = html.escape(str(scenarios.get("base", "N/A")))
            risk_val = html.escape(str(ebitda_projection.get("risk_adjusted", "N/A")))
            ebitda_bullets = ""
            if scenarios:
                for scenario, value in scenarios.items():
                    ebitda_bullets += f"<li>{html.escape(str(scenario).title())}: {html.escape(str(value))}</li>\n"
                ebitda_bullets = f"<ul>{ebitda_bullets}</ul>"

            # IC recommendation
            if org_air >= 75:
                rec_text = "PROCEED — strong AI readiness supports full capital deployment"
                rec_class = "rec-proceed"
            elif org_air >= 60:
                rec_text = "PROCEED WITH MONITORING — address top 2 dimension gaps within 90 days"
                rec_class = "rec-monitor"
            else:
                rec_text = "CONDITIONAL — address critical gaps before next capital deployment"
                rec_class = "rec-conditional"

            html_str = f"""\
<!DOCTYPE html>
<html>
<head><meta charset="utf-8"><style>{IC_MEMO_CSS}</style></head>
<body>
    <h1>Investment Committee Memo: {cid}</h1>
    <p class="meta">Date: {html.escape(date_str)}</p>
    <p class="meta">Prepared by: PE Org-AI-R Agentic Platform</p>
    <p class="confidential">CONFIDENTIAL</p>

    <h2>Executive Summary</h2>
    <p>{cid} has an Org-AI-R score of {org_air:.1f}, reflecting V^R of {vr:.1f}
       and H^R of {hr:.1f}. {html.escape(readiness)}</p>

    <h2>AI Readiness Assessment</h2>
    <table>
        <thead><tr><th>Metric</th><th>Score</th><th>Benchmark</th></tr></thead>
        <tbody>{scoring_rows}</tbody>
    </table>
    {dim_html}

    <h2>Gap Analysis &amp; 100-Day Plan</h2>
    {gaps_html}

    <h2>EBITDA Impact Projection</h2>
    <p>Base case EBITDA improvement: {base_val} | Risk-adjusted: {risk_val}</p>
    {ebitda_bullets}

    <h2>IC Recommendation</h2>
    <p class="recommendation {rec_class}">Recommendation: {html.escape(rec_text)}</p>
</body>
</html>"""

            path = output_path or f"ic_memo_{company_id}_{datetime.now().strftime('%Y%m%d')}.pdf"
            HTML(string=html_str).write_pdf(path)
            logger.info("IC memo saved to %s", path)
            return path
        except Exception as exc:
            logger.exception("IC memo PDF generation failed; falling back to text: %s", exc)
            return self._generate_text(
                company_id, org_air, vr, hr, gap_analysis, ebitda_projection, output_path, date_str
            )

    def _generate_text(
        self,
        company_id: str,
        org_air: float,
        vr: float,
        hr: float,
        gap_analysis: Dict[str, Any],
        ebitda_projection: Dict[str, Any],
        output_path: Optional[str],
        date_str: str,
    ) -> str:
        path = output_path or f"ic_memo_{company_id}_{datetime.now().strftime('%Y%m%d')}.txt"
        scenarios = ebitda_projection.get("scenarios", {})
        with open(path, "w", encoding="utf-8") as f:
            f.write(f"INVESTMENT COMMITTEE MEMO: {company_id}\n")
            f.write(f"Date: {date_str} | CONFIDENTIAL\n\n")
            f.write(f"Org-AI-R: {org_air:.1f} | V^R: {vr:.1f} | H^R: {hr:.1f}\n")
            f.write(f"EBITDA Base: {scenarios.get('base', 'N/A')} | "
                    f"Risk-adj: {ebitda_projection.get('risk_adjusted', 'N/A')}\n")
        return path

    def _generate_docx(
        self,
        Document: Any,
        company_id: str,
        org_air: float,
        vr: float,
        hr: float,
        dimension_scores: Dict[str, Any],
        gap_analysis: Dict[str, Any],
        ebitda_projection: Dict[str, Any],
        output_path: Optional[str],
        date_str: str,
    ) -> str:
        path = output_path or f"ic_memo_{company_id}_{datetime.now().strftime('%Y%m%d')}.docx"
        if not path.lower().endswith(".docx"):
            path = path + ".docx"

        doc = Document()
        doc.add_heading(f"Investment Committee Memo: {company_id}", level=1)
        doc.add_paragraph(f"Date: {date_str}")
        doc.add_paragraph("CONFIDENTIAL")

        doc.add_heading("Executive Summary", level=2)
        readiness = (
            "Strong AI readiness positions for value creation."
            if org_air >= 70
            else "Improvement opportunities identified across key dimensions."
        )
        doc.add_paragraph(
            f"{company_id} has an Org-AI-R score of {org_air:.1f}, reflecting V^R of {vr:.1f} "
            f"and H^R of {hr:.1f}. {readiness}"
        )

        doc.add_heading("AI Readiness Assessment", level=2)
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "Metric"
        hdr[1].text = "Score"
        hdr[2].text = "Benchmark"
        for metric, score, bench in [
            ("Org-AI-R", org_air, 65.0),
            ("V^R (Valuation Readiness)", vr, 60.0),
            ("H^R (Human Capital Risk)", hr, 60.0),
        ]:
            row = table.add_row().cells
            row[0].text = str(metric)
            row[1].text = f"{float(score):.1f}"
            row[2].text = f"{float(bench):.1f}"

        if dimension_scores:
            doc.add_heading("Dimension Scores", level=3)
            dim_table = doc.add_table(rows=1, cols=2)
            dim_hdr = dim_table.rows[0].cells
            dim_hdr[0].text = "Dimension"
            dim_hdr[1].text = "Score"
            for dim, score in dimension_scores.items():
                row = dim_table.add_row().cells
                row[0].text = str(dim).replace("_", " ").title()
                try:
                    row[1].text = f"{float(score or 0.0):.1f}"
                except Exception:
                    row[1].text = str(score)

        doc.add_heading("Gap Analysis & 100-Day Plan", level=2)
        gaps = gap_analysis.get("dimension_gaps", []) or []
        if gaps:
            for gap in gaps[:8]:
                dim_name = str(gap.get("dimension", "")).replace("_", " ").title()
                current = gap.get("current_score", "N/A")
                target = gap.get("target_score", "N/A")
                priority = gap.get("priority", "N/A")
                doc.add_paragraph(
                    f"{dim_name}: current {current} → target {target} [Priority: {priority}]",
                    style="List Bullet",
                )
        else:
            doc.add_paragraph("No gap analysis data available.")

        doc.add_heading("EBITDA Impact Projection", level=2)
        scenarios = ebitda_projection.get("scenarios", {}) or {}
        doc.add_paragraph(
            f"Risk-adjusted: {ebitda_projection.get('risk_adjusted', 'N/A')} | "
            f"Delta Org-AI-R: {ebitda_projection.get('delta_air', 'N/A')}"
        )
        if scenarios:
            for scenario, value in scenarios.items():
                doc.add_paragraph(f"{str(scenario).title()}: {value}", style="List Bullet")

        doc.add_heading("IC Recommendation", level=2)
        if org_air >= 75:
            rec_text = "PROCEED — strong AI readiness supports full capital deployment"
        elif org_air >= 60:
            rec_text = "PROCEED WITH MONITORING — address top 2 dimension gaps within 90 days"
        else:
            rec_text = "CONDITIONAL — address critical gaps before next capital deployment"
        doc.add_paragraph(rec_text)

        doc.save(path)
        logger.info("IC memo saved to %s", path)
        return path


# Module-level singleton
ic_memo_generator = ICMemoGenerator()
