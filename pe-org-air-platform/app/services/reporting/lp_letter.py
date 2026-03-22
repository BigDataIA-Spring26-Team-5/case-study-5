"""LP Letter Generator — produces LP update letter summarizing Fund-AI-R metrics.

Requires:  python-docx>=1.1.0  (add to requirements.txt)
"""
from __future__ import annotations

import logging
from datetime import datetime
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)

LP_LETTER_TEMPLATE = """\
Dear Limited Partners,

We are pleased to share the quarterly AI Readiness update for {fund_id}.

FUND PERFORMANCE SUMMARY
-------------------------
Fund-AI-R Score:     {fund_air:.1f} / 100
Portfolio Companies: {company_count}
AI Leaders (≥70):    {leaders}
AI Laggards (<50):   {laggards}

PORTFOLIO HIGHLIGHTS
--------------------
{highlights}

VALUE CREATION PIPELINE
------------------------
{value_creation}

Our agentic due diligence platform continues to surface actionable insights
across the portfolio. The Org-AI-R framework identifies specific improvement
vectors for each company, enabling precise capital deployment decisions.

We remain committed to driving AI-readiness across the portfolio and look
forward to sharing further progress in our next update.

Sincerely,
The Investment Team
{fund_id}
{date}
"""


class LPLetterGenerator:
    """Generates LP update letters from Fund-AI-R metrics.

    Falls back to plain-text if python-docx is not installed.
    """

    def _try_docx(self) -> Optional[Any]:
        try:
            from docx import Document  # type: ignore[import]
            return Document
        except ImportError:
            return None

    def generate(
        self,
        fund_id: str,
        fund_metrics: Dict[str, Any],
        company_scores: List[Dict[str, Any]],
        output_path: Optional[str] = None,
    ) -> str:
        """Generate LP update letter and save as .docx (or .txt fallback).

        Args:
            fund_id: Fund identifier
            fund_metrics: Dict with fund_air, company_count, etc.
            company_scores: List of dicts with ticker, org_air, sector
            output_path: Override default file path

        Returns:
            Path to the generated file.
        """
        leaders = sum(1 for c in company_scores if c.get("org_air", 0) >= 70)
        laggards = sum(1 for c in company_scores if c.get("org_air", 0) < 50)
        date_str = datetime.utcnow().strftime("%B %d, %Y")

        sorted_companies = sorted(
            company_scores, key=lambda x: x.get("org_air", 0), reverse=True
        )
        highlights = "\n".join([
            f"  • {c.get('ticker', '?')}: Org-AI-R {c.get('org_air', 0):.1f} "
            f"({'Leader' if c.get('org_air', 0) >= 70 else 'Developing'})"
            for c in sorted_companies
        ])

        top = sorted_companies[0] if sorted_companies else {}
        value_creation = (
            f"  • {top.get('ticker', 'N/A')} leads the portfolio at "
            f"{top.get('org_air', 0):.1f} Org-AI-R. "
            f"Gap analysis identifies data infrastructure and talent as primary value levers."
        )

        # Get fund_air from the right field name (handles both CS5 and legacy)
        fund_air = fund_metrics.get("fund_air", fund_metrics.get("fund_air_score", 0.0))

        letter = LP_LETTER_TEMPLATE.format(
            fund_id=fund_id,
            fund_air=fund_air,
            company_count=len(company_scores),
            leaders=leaders,
            laggards=laggards,
            highlights=highlights,
            value_creation=value_creation,
            date=date_str,
        )

        Document = self._try_docx()
        path = output_path or f"lp_letter_{fund_id}_{datetime.now().strftime('%Y%m%d')}.docx"

        if Document is None:
            logger.warning("python-docx not installed — saving as .txt")
            txt_path = path.replace(".docx", ".txt")
            with open(txt_path, "w") as f:
                f.write(letter)
            logger.info("LP letter saved to %s", txt_path)
            return txt_path

        doc = Document()
        doc.add_heading(f"LP Update — {fund_id}", 0)
        for line in letter.strip().split("\n"):
            if line.startswith("---"):
                doc.add_paragraph("─" * 40)
            elif line.isupper() and len(line) > 5:
                doc.add_heading(line.title(), level=2)
            else:
                doc.add_paragraph(line)

        doc.save(path)
        logger.info("LP letter saved to %s", path)
        return path


# Module-level singleton
lp_letter_generator = LPLetterGenerator()
